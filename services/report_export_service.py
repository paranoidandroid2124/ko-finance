"""Utilities to export investment memos as PDF, Word, or Excel documents."""

from __future__ import annotations

import base64
import csv
import io
import logging
import re
from typing import Any, Dict, Iterable, List, Optional

import fitz
import markdown
from docx import Document
from openpyxl import Workbook
from openpyxl.utils import get_column_letter

try:  # pragma: no cover - optional pandas dependency
    import pandas as pd  # type: ignore
except ImportError:  # pragma: no cover
    pd = None

PDF_PAGE_RECT = fitz.Rect(36, 36, 559, 806)  # letter-ish margins


def _build_html_document(title: str, markdown_body: str, sources: Iterable[Dict[str, str]]) -> str:
    body_html = markdown.markdown(
        markdown_body,
        extensions=[
            "tables",
            "fenced_code",
            "toc",
        ],
    )
    references = ""
    normalized_sources = list(sources)
    if normalized_sources:
        refs = "".join(
            f"<li><a href='{entry.get('url','')}'>{entry.get('title','Source')}</a> "
            f"<em>{entry.get('date','')}</em></li>"
            for entry in normalized_sources
        )
        references = f"<h3>References</h3><ol>{refs}</ol>"
    return f"""
        <html>
            <head>
                <meta charset="utf-8" />
                <style>
                    body {{
                        font-family: 'Helvetica', 'Arial', sans-serif;
                        font-size: 12pt;
                        line-height: 1.6;
                        color: #111827;
                    }}
                    h1, h2, h3 {{
                        color: #0f172a;
                    }}
                    ol {{
                        padding-left: 18px;
                    }}
                    a {{
                        color: #1d4ed8;
                        text-decoration: none;
                    }}
                </style>
            </head>
            <body>
                <h1>{title}</h1>
                {body_html}
                {references}
            </body>
        </html>
    """


def export_pdf(
    *,
    title: str,
    markdown_body: str,
    sources: Iterable[Dict[str, str]],
    chart_image: Optional[str] = None,
) -> bytes:
    """Render Markdown content into a PDF byte stream."""

    html_document = _build_html_document(title, markdown_body, sources)
    doc = fitz.open()
    page = doc.new_page()
    page.insert_htmlbox(PDF_PAGE_RECT, html_document)
    if chart_image:
        try:
            _, _, payload = chart_image.partition(",")
            image_bytes = base64.b64decode(payload or chart_image)
            chart_page = doc.new_page()
            chart_page.insert_image(PDF_PAGE_RECT, stream=image_bytes)
        except Exception:
            pass
    pdf_bytes = doc.tobytes()
    doc.close()
    return pdf_bytes


def export_docx(
    *,
    title: str,
    markdown_body: str,
    sources: Iterable[Dict[str, str]],
    chart_image: Optional[str] = None,
) -> bytes:
    """Render Markdown-like content into a .docx byte stream."""

    document = Document()
    document.add_heading(title, level=0)

    if chart_image:
        try:
            _, _, payload = chart_image.partition(",")
            image_bytes = base64.b64decode(payload or chart_image)
            image_stream = io.BytesIO(image_bytes)
            document.add_heading("Performance Chart", level=1)
            document.add_picture(image_stream, width=document.sections[0].page_width - document.sections[0].left_margin - document.sections[0].right_margin)
        except Exception:
            pass

    for line in markdown_body.splitlines():
        stripped = line.strip()
        if not stripped:
            document.add_paragraph("")
            continue
        if stripped.startswith("# "):
            document.add_heading(stripped[2:].strip(), level=1)
        elif stripped.startswith("## "):
            document.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith("- "):
            document.add_paragraph(stripped[2:].strip(), style="List Bullet")
        else:
            document.add_paragraph(stripped)

    normalized_sources: List[Dict[str, str]] = list(sources)
    if normalized_sources:
        document.add_page_break()
        document.add_heading("References", level=1)
        for entry in normalized_sources:
            title_text = entry.get("title") or entry.get("url") or "Source"
            paragraph = document.add_paragraph(style="List Number")
            run = paragraph.add_run(f"{title_text} ")
            run.bold = True
            date_str = entry.get("date")
            if date_str:
                paragraph.add_run(f"({date_str}) ")
            url = entry.get("url")
            if url:
                paragraph.add_run(url)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


logger = logging.getLogger(__name__)


def _sanitize_sheet_name(label: str, suffix: Optional[int] = None) -> str:
    base = re.sub(r"[\\/*?:\[\]]", "", label).strip() or "Table"
    truncated = base[:25]
    if suffix is not None:
        return f"{truncated}_{suffix}"
    return truncated


def _normalize_markdown_line(line: str) -> str:
    stripped = line.strip()
    stripped = stripped.lstrip("#").lstrip("-*>").strip()
    return stripped


def _parse_markdown_table_block(block_lines: List[str], title: Optional[str]) -> Optional[Dict[str, Any]]:
    block_text = "\n".join(block_lines).strip()
    if not block_text:
        return None
    try:
        meaningful_lines: List[str] = []
        for line in block_lines:
            stripped = line.strip()
            if not stripped:
                continue
            if set(stripped) <= {"|", "-", ":", " "}:
                continue
            meaningful_lines.append(stripped)
        if not meaningful_lines:
            return _fallback_table(title, block_text)
        buffer = io.StringIO("\n".join(meaningful_lines))
        reader = csv.reader(buffer, delimiter="|", quotechar='"')
        data: List[List[str]] = []
        for row in reader:
            cleaned_row = [cell.strip() for cell in row if cell.strip() != ""]
            if cleaned_row:
                data.append(cleaned_row)
        if len(data) < 2:
            return _fallback_table(title, block_text)
        headers = data[0]
        rows = data[1:]
        width = max(len(headers), *(len(row) for row in rows)) if rows else len(headers)
        if width == 0:
            return _fallback_table(title, block_text)
        normalized_headers = headers + [f"Column {idx + 1}" for idx in range(len(headers), width)]
        normalized_headers = normalized_headers[:width]
        normalized_rows: List[List[str]] = []
        for row in rows:
            padded = row + [""] * (width - len(row))
            normalized_rows.append(padded[:width])
        return {
            "title": title or "Table",
            "headers": normalized_headers,
            "rows": normalized_rows,
        }
    except Exception as exc:  # pragma: no cover - best effort
        logger.warning("Markdown table parsing failed; falling back to raw text: %s", exc, exc_info=True)
        return _fallback_table(title, block_text)


def _fallback_table(title: Optional[str], block_text: str) -> Dict[str, Any]:
    return {
        "title": title or "Table",
        "headers": ["Raw Table"],
        "rows": [[block_text]],
    }


def _extract_markdown_tables(markdown_body: str) -> List[Dict[str, Any]]:
    tables: List[Dict[str, Any]] = []
    lines = markdown_body.splitlines()
    buffer: List[str] = []
    last_heading: Optional[str] = None
    for line in lines:
        stripped = line.strip()
        if stripped.startswith("#"):
            last_heading = stripped.lstrip("#").strip()
        if stripped.startswith("|") and stripped.count("|") >= 2:
            buffer.append(stripped)
            continue
        if buffer:
            table = _parse_markdown_table_block(buffer, last_heading)
            if table:
                tables.append(table)
            buffer = []
    if buffer:
        table = _parse_markdown_table_block(buffer, last_heading)
        if table:
            tables.append(table)
    return tables


def _rows_to_dataframe(headers: List[str], rows: List[List[str]]):
    if pd is None:
        return None
    normalized_headers = [header or f"Column {idx + 1}" for idx, header in enumerate(headers)]
    normalized_rows: List[Dict[str, Any]] = []
    for row in rows:
        record: Dict[str, Any] = {}
        for idx, header in enumerate(normalized_headers):
            record[header] = row[idx] if idx < len(row) else ""
        normalized_rows.append(record)
    return pd.DataFrame(normalized_rows)


def _auto_fit_columns(ws) -> None:
    for column_cells in ws.columns:
        max_length = 0
        column = column_cells[0].column if column_cells else 1
        for cell in column_cells:
            if cell.value:
                max_length = max(max_length, len(str(cell.value)))
        adjusted = max_length + 2
        ws.column_dimensions[get_column_letter(column)].width = min(adjusted, 60)


def export_excel(
    *,
    title: str,
    markdown_body: str,
    sources: Iterable[Dict[str, str]],
    key_stats: Optional[Dict[str, Any]] = None,
) -> bytes:
    """Render memo content and extracted tables into an Excel workbook."""

    workbook = Workbook()
    summary_ws = workbook.active
    summary_ws.title = "Summary"
    summary_ws.append([title])
    summary_ws.append([])
    for line in markdown_body.splitlines():
        normalized = _normalize_markdown_line(line)
        if not normalized:
            continue
        summary_ws.append([normalized])
    _auto_fit_columns(summary_ws)

    stats_ws = workbook.create_sheet("Key Stats")
    if key_stats:
        stats_ws.append(["Ticker", key_stats.get("ticker")])
        stats_ws.append(["Event Type", key_stats.get("eventType")])
        stats_ws.append(["Event Date", key_stats.get("eventDate")])
        window = key_stats.get("window") or {}
        if window:
            stats_ws.append(["Window", window.get("label") or f"[{window.get('start')},{window.get('end')}]"])
            stats_ws.append(["Significance", window.get("significance")])
        metrics = key_stats.get("metrics") or {}
        if metrics:
            stats_ws.append([])
            stats_ws.append(["Sample Size", metrics.get("sampleSize")])
            stats_ws.append(["Mean CAAR", metrics.get("meanCaar")])
            stats_ws.append(["Hit Rate", metrics.get("hitRate")])
            stats_ws.append(["CI Low", metrics.get("ciLo")])
            stats_ws.append(["CI High", metrics.get("ciHi")])
            stats_ws.append(["P-Value", metrics.get("pValue")])
        recent_section = key_stats.get("recentEvents") or {}
        recent_events = []
        if isinstance(recent_section, dict):
            recent_events = recent_section.get("events") or []
        elif isinstance(recent_section, list):
            recent_events = recent_section
        if recent_events:
            stats_ws.append([])
            stats_ws.append(["Recent Events"])
            stats_ws.append(["Date", "Title", "CAAR"])
            for entry in recent_events:
                stats_ws.append(
                    [
                        entry.get("event_date") or entry.get("eventDate") or entry.get("published_at"),
                        entry.get("title") or entry.get("corp_name"),
                        entry.get("caar"),
                    ]
                )
    else:
        stats_ws.append(["No key stats provided."])
    _auto_fit_columns(stats_ws)

    tables = _extract_markdown_tables(markdown_body)
    for idx, table in enumerate(tables, start=1):
        sheet_name = _sanitize_sheet_name(table.get("title") or f"Table {idx}", suffix=idx)
        ws = workbook.create_sheet(sheet_name)
        headers = table.get("headers") or []
        rows = table.get("rows") or []
        if headers:
            ws.append(headers)
        for row in rows:
            if headers and len(row) < len(headers):
                row = row + [""] * (len(headers) - len(row))
            ws.append(row)
        _auto_fit_columns(ws)

    sources_ws = workbook.create_sheet("Sources")
    sources_ws.append(["Title", "URL", "Date"])
    for entry in sources:
        sources_ws.append([entry.get("title"), entry.get("url"), entry.get("date")])
    _auto_fit_columns(sources_ws)

    buffer = io.BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


__all__ = ["export_pdf", "export_docx", "export_excel"]
